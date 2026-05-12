"""
blueprints/cv.py — CV Builder Blueprint
=========================================
Handles CV generation in PDF and DOCX formats for
different countries/regions with 4 template types.

Routes (prefix: /cv-builder):
    GET  /                        → Country selection page
    GET  /cv_form/<country_id>    → CV form for selected country
    POST /generate_cv             → Generate and download CV file

Template Types:
    ats           → USA, Canada, Australia, Ireland (no photo)
    photo_personal → Germany, Austria, Turkey (photo + personal info)
    passport      → China, UAE, Japan (passport photo + number)
    europass      → Belgium, Denmark, Italy, France (European standard)
"""

from flask import Blueprint, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from io import BytesIO
from PIL import Image
import base64

# Blueprint with /cv-builder prefix (registered in app.py)
cv_bp = Blueprint('cv', __name__, url_prefix='/cv-builder')


# ==============================================================
#   Template Type Configurations
# ==============================================================

TEMPLATE_CONFIGS = {
    'ats': {
        'name':                    'ATS-Friendly',
        'display_name':            'USA, Canada, Australia, Ireland',
        'requires_photo':          False,
        'requires_personal_details': False,
        'max_pages':               2,
        'description':             'No photo needed, ATS-optimized, 1-2 pages'
    },
    'photo_personal': {
        'name':                    'Photo + Personal Details',
        'display_name':            'Germany, Austria, Turkey',
        'requires_photo':          True,
        'requires_personal_details': True,
        'max_pages':               3,
        'description':             'Professional photo required, include DOB, nationality, marital status'
    },
    'passport': {
        'name':                    'Passport/Visa Details',
        'display_name':            'China, UAE, Japan',
        'requires_photo':          True,
        'requires_personal_details': True,
        'max_pages':               2,
        'description':             'Passport photo and number required, visa status for UAE'
    },
    'europass': {
        'name':                    'Europass Style',
        'display_name':            'Belgium, Denmark, Italy, France',
        'requires_photo':          False,
        'requires_personal_details': False,
        'max_pages':               2,
        'description':             'Languages section very important, European standard format'
    }
}


# ==============================================================
#   Country → Template Mapping
# ==============================================================

COUNTRY_TEMPLATE = {
    'usa':       {'name': 'USA',       'flag': '🇺🇸', 'template': 'ats',            'group': 'A'},
    'canada':    {'name': 'Canada',    'flag': '🇨🇦', 'template': 'ats',            'group': 'A'},
    'australia': {'name': 'Australia', 'flag': '🇦🇺', 'template': 'ats',            'group': 'A'},
    'ireland':   {'name': 'Ireland',   'flag': '🇮🇪', 'template': 'ats',            'group': 'A'},
    'germany':   {'name': 'Germany',   'flag': '🇩🇪', 'template': 'photo_personal', 'group': 'B'},
    'austria':   {'name': 'Austria',   'flag': '🇦🇹', 'template': 'photo_personal', 'group': 'B'},
    'turkey':    {'name': 'Turkey',    'flag': '🇹🇷', 'template': 'photo_personal', 'group': 'B'},
    'china':     {'name': 'China',     'flag': '🇨🇳', 'template': 'passport',       'group': 'C'},
    'uae':       {'name': 'UAE',       'flag': '🇦🇪', 'template': 'passport',       'group': 'C'},
    'japan':     {'name': 'Japan',     'flag': '🇯🇵', 'template': 'passport',       'group': 'C'},
    'belgium':   {'name': 'Belgium',   'flag': '🇧🇪', 'template': 'europass',       'group': 'D'},
    'denmark':   {'name': 'Denmark',   'flag': '🇩🇰', 'template': 'europass',       'group': 'D'},
    'italy':     {'name': 'Italy',     'flag': '🇮🇹', 'template': 'europass',       'group': 'D'},
    'france':    {'name': 'France',    'flag': '🇫🇷', 'template': 'europass',       'group': 'D'},
}


# ==============================================================
#   Routes
# ==============================================================

@cv_bp.route('/')
def cv_index():
    """CV Builder home — country selection page."""
    return render_template(
        'cv-builder.html',
        countries        = COUNTRY_TEMPLATE,
        template_configs = TEMPLATE_CONFIGS
    )


@cv_bp.route('/cv_form/<country_id>')
def cv_form(country_id):
    """
    CV input form for a specific country.

    Args:
        country_id: e.g. 'usa', 'germany', 'china'
    """
    if country_id not in COUNTRY_TEMPLATE:
        return "Country not found", 404

    country         = COUNTRY_TEMPLATE[country_id]
    template_config = TEMPLATE_CONFIGS[country['template']]

    return render_template(
        'cv_form.html',
        country         = country,
        template_config = template_config,
        template_type   = country['template']
    )


@cv_bp.route('/generate_cv', methods=['POST'])
def generate_cv():
    """
    Generate and return a CV file (PDF or DOCX).

    Expects JSON:
        {
            country_id:    str,
            template_type: str,
            cv_data:       dict,
            format:        'pdf' | 'docx',
            photo_data:    str (base64, optional)
        }
    """
    data          = request.json or {}
    country_id    = data.get('country_id', 'usa')
    template_type = data.get('template_type', 'ats')
    cv_data       = data.get('cv_data', {})
    format_type   = data.get('format', 'pdf')
    photo_data    = data.get('photo_data', None)   # Base64 encoded image

    if format_type == 'pdf':
        return _generate_pdf(template_type, cv_data, country_id, photo_data)
    else:
        return _generate_docx(template_type, cv_data, country_id, photo_data)


# ==============================================================
#   Image Helper Functions
# ==============================================================

def _resize_image_for_pdf(image_data: str, max_width: int = 80, max_height: int = 100):
    """
    Decode, resize, and convert image for embedding in PDF.
    Converts RGBA → RGB since PDF doesn't support transparency.

    Args:
        image_data:  Base64 image string (with or without data URI prefix).
        max_width:   Maximum width in pixels.
        max_height:  Maximum height in pixels.

    Returns:
        BytesIO JPEG image, or None on failure.
    """
    try:
        # Strip data URI prefix if present
        if ',' in image_data:
            image_data = image_data.split(',')[1]

        image_bytes = base64.b64decode(image_data)
        img         = Image.open(BytesIO(image_bytes))

        # Convert RGBA → RGB (PDFs don't support alpha channel)
        if img.mode == 'RGBA':
            background = Image.new('RGB', img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])
            img = background

        img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)

        output = BytesIO()
        img.save(output, format='JPEG', quality=85)
        output.seek(0)
        return output

    except Exception as e:
        print(f"❌ PDF image resize error: {e}")
        return None


def _resize_image_for_docx(image_data: str):
    """
    Decode and convert image for embedding in DOCX.

    Args:
        image_data: Base64 image string.

    Returns:
        BytesIO JPEG image, or None on failure.
    """
    try:
        if ',' in image_data:
            image_data = image_data.split(',')[1]

        image_bytes = base64.b64decode(image_data)
        img         = Image.open(BytesIO(image_bytes))

        if img.mode == 'RGBA':
            background = Image.new('RGB', img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])
            img = background

        output = BytesIO()
        img.save(output, format='JPEG', quality=85)
        output.seek(0)
        return output

    except Exception as e:
        print(f"❌ DOCX image resize error: {e}")
        return None


# ==============================================================
#   PDF Generation
# ==============================================================

def _generate_pdf(template_type: str, cv_data: dict, country_id: str, photo_data):
    """
    Build and return a PDF CV using ReportLab.

    Args:
        template_type: One of the TEMPLATE_CONFIGS keys.
        cv_data:       Dictionary of CV field values from the form.
        country_id:    Used in the filename.
        photo_data:    Base64 photo string or None.
    """
    buffer = BytesIO()
    c      = canvas.Canvas(buffer, pagesize=letter)
    config = TEMPLATE_CONFIGS.get(template_type, {})
    y      = 750   # Start position from top of page

    # ── Name ─────────────────────────────────────────────────
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, cv_data.get('full_name', 'NAME NOT PROVIDED'))

    # ── Contact Info ─────────────────────────────────────────
    y -= 30
    c.setFont("Helvetica", 10)
    c.drawString(50,  y, f"Email: {cv_data.get('email', '')}")
    c.drawString(250, y, f"Phone: {cv_data.get('phone', '')}")
    if cv_data.get('linkedin'):
        y -= 15
        c.drawString(50, y, f"LinkedIn: {cv_data.get('linkedin', '')}")
    if cv_data.get('github'):
        y -= 15
        c.drawString(50, y, f"GitHub: {cv_data.get('github', '')}")

    # ── Photo (for templates that require it) ────────────────
    if config.get('requires_photo') and photo_data:
        try:
            img_bytes = _resize_image_for_pdf(photo_data)
            if img_bytes:
                c.drawImage(ImageReader(img_bytes), 450, 700, width=80, height=100, preserveAspectRatio=True)
        except Exception as e:
            print(f"❌ PDF photo embed error: {e}")
            c.drawString(450, 745, "[Photo error]")
    elif config.get('requires_photo'):
        c.drawString(450, 745, "[Photo not provided]")

    # ── Personal Details (Germany/China/UAE templates) ────────
    if config.get('requires_personal_details'):
        personal_fields = [
            ('dob',            'Date of Birth'),
            ('nationality',    'Nationality'),
            ('marital_status', 'Marital Status'),
            ('passport_no',    'Passport No'),
            ('visa_status',    'Visa Status'),
        ]
        if any(cv_data.get(f) for f, _ in personal_fields):
            y -= 50
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y, "PERSONAL DETAILS")
            c.line(50, y - 5, 500, y - 5)
            y -= 25
            c.setFont("Helvetica", 10)
            for field, label in personal_fields:
                if cv_data.get(field):
                    c.drawString(50, y, f"{label}: {cv_data[field]}")
                    y -= 20

    # ── Education ─────────────────────────────────────────────
    if cv_data.get('degree') or cv_data.get('university'):
        y -= 40
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, "EDUCATION")
        c.line(50, y - 5, 500, y - 5)
        y -= 25
        c.setFont("Helvetica", 10)
        c.drawString(50,  y, cv_data.get('degree', ''))
        c.drawString(300, y, cv_data.get('university', ''))
        if cv_data.get('gpa') or cv_data.get('graduation_year'):
            y -= 20
            c.drawString(50,  y, f"GPA: {cv_data.get('gpa', '')}")
            c.drawString(300, y, f"Expected: {cv_data.get('graduation_year', '')}")

    # ── Work Experience ───────────────────────────────────────
    if cv_data.get('work_experience') and y > 150:
        y -= 40
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, "WORK EXPERIENCE")
        c.line(50, y - 5, 500, y - 5)
        y -= 25
        c.setFont("Helvetica", 10)
        for line in cv_data['work_experience'].split('\n')[:4]:
            if line.strip():
                c.drawString(50, y, f"• {line[:80]}")
                y -= 20
                if y < 100:
                    c.showPage(); y = 750

    # ── Projects ──────────────────────────────────────────────
    if cv_data.get('projects') and y > 150:
        y -= 40
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, "PROJECTS")
        c.line(50, y - 5, 500, y - 5)
        y -= 25
        c.setFont("Helvetica", 10)
        for line in cv_data['projects'].split('\n')[:4]:
            if line.strip():
                c.drawString(50, y, f"• {line.lstrip('•-* ').strip()[:75]}")
                y -= 20
                if y < 100:
                    c.showPage(); y = 750

    # ── Skills ────────────────────────────────────────────────
    if cv_data.get('skills') and y > 120:
        y -= 40
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, "SKILLS")
        c.line(50, y - 5, 500, y - 5)
        y -= 25
        c.setFont("Helvetica", 10)
        c.drawString(50, y, cv_data['skills'][:100])

    # ── Languages ─────────────────────────────────────────────
    if cv_data.get('languages') and y > 100:
        y -= 40
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, "LANGUAGES")
        c.line(50, y - 5, 500, y - 5)
        y -= 25
        c.setFont("Helvetica", 10)
        c.drawString(50, y, cv_data['languages'][:100])

    # ── Certifications ────────────────────────────────────────
    if cv_data.get('certifications') and y > 100:
        y -= 40
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, "CERTIFICATIONS")
        c.line(50, y - 5, 500, y - 5)
        y -= 25
        c.setFont("Helvetica", 10)
        c.drawString(50, y, cv_data['certifications'][:100])

    c.save()
    buffer.seek(0)

    filename = f"CV_{country_id}_{cv_data.get('full_name', 'scholar').replace(' ', '_')}.pdf"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/pdf')


# ==============================================================
#   DOCX Generation
# ==============================================================

def _generate_docx(template_type: str, cv_data: dict, country_id: str, photo_data):
    """
    Build and return a DOCX CV using python-docx.

    Args:
        template_type: One of the TEMPLATE_CONFIGS keys.
        cv_data:       Dictionary of CV field values from the form.
        country_id:    Used in the filename.
        photo_data:    Base64 photo string or None.
    """
    doc    = Document()
    config = TEMPLATE_CONFIGS.get(template_type, {})

    # ── Header: Photo Template (side-by-side table) ───────────
    if config.get('requires_photo'):
        table               = doc.add_table(rows=1, cols=2)
        table.autofit       = False
        table.columns[0].width = Cm(12)
        table.columns[1].width = Cm(4)

        # Left cell: name + contact
        left     = table.cell(0, 0)
        name_run = left.paragraphs[0].add_run(cv_data.get('full_name', 'Curriculum Vitae'))
        name_run.bold            = True
        name_run.font.size       = Pt(16)
        contact                  = left.add_paragraph()
        contact.add_run(f"Email: {cv_data.get('email', '')} | Phone: {cv_data.get('phone', '')}")

        # Right cell: photo
        right = table.cell(0, 1)
        if photo_data:
            try:
                img_bytes  = _resize_image_for_docx(photo_data)
                if img_bytes:
                    photo_para           = right.add_paragraph()
                    photo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    photo_para.add_run().add_picture(img_bytes, width=Cm(3.5), height=Cm(4.5))
            except Exception as e:
                print(f"❌ DOCX photo embed error: {e}")
                right.add_paragraph("[Photo error]")
        doc.add_paragraph()

    else:
        # ── Header: No-photo template ────────────────────────
        doc.add_heading(cv_data.get('full_name', 'Curriculum Vitae'), 0)
        doc.add_paragraph(f"Email: {cv_data.get('email', '')} | Phone: {cv_data.get('phone', '')}")
        if cv_data.get('linkedin'): doc.add_paragraph(f"LinkedIn: {cv_data['linkedin']}")
        if cv_data.get('github'):   doc.add_paragraph(f"GitHub: {cv_data['github']}")
        doc.add_paragraph()

    # ── Personal Details ──────────────────────────────────────
    if config.get('requires_personal_details'):
        personal_fields = {
            'dob':            'Date of Birth',
            'nationality':    'Nationality',
            'marital_status': 'Marital Status',
            'passport_no':    'Passport Number',
            'visa_status':    'Visa Status',
        }
        personal_lines = [
            f"{label}: {cv_data[key]}"
            for key, label in personal_fields.items() if cv_data.get(key)
        ]
        if personal_lines:
            doc.add_heading('PERSONAL INFORMATION', level=1)
            for line in personal_lines:
                doc.add_paragraph(line)
            doc.add_paragraph()

    # ── Education ─────────────────────────────────────────────
    if cv_data.get('degree') or cv_data.get('university'):
        doc.add_heading('EDUCATION', level=1)
        doc.add_paragraph(f"{cv_data.get('degree', '')} — {cv_data.get('university', '')}")
        if cv_data.get('gpa') or cv_data.get('graduation_year'):
            doc.add_paragraph(f"GPA: {cv_data.get('gpa', 'N/A')} | Expected: {cv_data.get('graduation_year', 'N/A')}")
        if cv_data.get('coursework'):
            doc.add_paragraph(f"Relevant Coursework: {cv_data['coursework']}")
        doc.add_paragraph()

    # ── Work Experience ───────────────────────────────────────
    if cv_data.get('work_experience'):
        doc.add_heading('WORK EXPERIENCE', level=1)
        for line in cv_data['work_experience'].split('\n'):
            if line.strip():
                doc.add_paragraph(f"• {line.strip()}", style='List Bullet')
        doc.add_paragraph()

    # ── Projects ──────────────────────────────────────────────
    if cv_data.get('projects'):
        doc.add_heading('PROJECTS', level=1)
        for line in cv_data['projects'].split('\n'):
            if line.strip():
                doc.add_paragraph(f"• {line.lstrip('•-* ').strip()}", style='List Bullet')
        doc.add_paragraph()

    # ── Skills ────────────────────────────────────────────────
    if cv_data.get('skills'):
        doc.add_heading('SKILLS', level=1)
        doc.add_paragraph(cv_data['skills'])
        doc.add_paragraph()

    # ── Languages ─────────────────────────────────────────────
    if cv_data.get('languages'):
        doc.add_heading('LANGUAGES', level=1)
        doc.add_paragraph(cv_data['languages'])
        doc.add_paragraph()

    # ── Certifications ────────────────────────────────────────
    if cv_data.get('certifications'):
        doc.add_heading('CERTIFICATIONS', level=1)
        doc.add_paragraph(cv_data['certifications'])
        doc.add_paragraph()

    # ── Publications ──────────────────────────────────────────
    if cv_data.get('publications'):
        doc.add_heading('PUBLICATIONS', level=1)
        doc.add_paragraph(cv_data['publications'])

    # ── Save & Return ─────────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"CV_{country_id}_{cv_data.get('full_name', 'scholar').replace(' ', '_')}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
