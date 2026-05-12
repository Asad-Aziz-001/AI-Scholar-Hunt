"""
essay.py — AI Essay Analyzer & Rewriter Blueprint
====================================================
Complete Pipeline Architecture:
  User Input → Preprocessing → NLP Engine →
  [Grammar, Vocabulary, Coherence, Readability, Tone,
   Redundancy, Structure, Arguments, Plagiarism] →
  LLM Rewriter → Scoring Engine → Feedback Generator →
  Output (Improved Essay + Score + Suggestions + Download)

Routes:
  GET  /essay                    → Essay page
  POST /api/essay/analyze        → Full analysis pipeline
  POST /api/essay/download/pdf   → Download as PDF
  POST /api/essay/download/docx  → Download as DOCX
"""

import re
import os
import json
import math
import string
import urllib.request
from io import BytesIO
from collections import Counter
from flask import Blueprint, render_template, request, jsonify, send_file
from flask_login import login_required

essay_bp = Blueprint('essay', __name__)

# ── Groq API Key ─────────────────────────────────────────
from dotenv import load_dotenv
from groq import Groq

load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
MODEL = os.getenv("MODEL", "mixtral-8x7b-32768")

if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY not found")

MAX_TOKENS = 1500
MAX_CONTEXT_CHARS = 5000

client = Groq(api_key=GROQ_API_KEY)


# ==============================================================
#   STEP 1 — PREPROCESSING
# ==============================================================

def preprocess_text(text: str) -> dict:
    """
    Tokenization, cleaning, and basic NLP stats.

    Returns:
        sentences, words, tokens, cleaned_text, word_freq
    """
    # Clean text
    cleaned = re.sub(r'\s+', ' ', text.strip())

    # Sentence tokenization
    sentences = re.split(r'(?<=[.!?])\s+', cleaned)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 5]

    # Word tokenization
    words = re.findall(r'\b[a-zA-Z]+\b', cleaned.lower())

    # Remove stopwords for analysis
    stopwords = {
        'the','a','an','and','or','but','in','on','at','to','for','of',
        'with','by','from','is','are','was','were','be','been','being',
        'have','has','had','do','does','did','will','would','could',
        'should','may','might','shall','this','that','these','those',
        'i','we','you','he','she','it','they','my','your','our','their'
    }
    content_words = [w for w in words if w not in stopwords and len(w) > 2]
    word_freq     = Counter(content_words)

    return {
        'sentences':    sentences,
        'words':        words,
        'content_words': content_words,
        'word_freq':    word_freq,
        'cleaned_text': cleaned,
        'word_count':   len(words),
        'sent_count':   len(sentences),
        'char_count':   len(cleaned)
    }


# ==============================================================
#   STEP 2 — NLP ENGINE (POS + NER + Patterns)
# ==============================================================

def run_nlp_engine(preprocessed: dict, text: str) -> dict:
    """
    Lightweight NLP without heavy dependencies.
    Detects POS patterns, named entities, and linguistic features.
    """
    words     = preprocessed['words']
    sentences = preprocessed['sentences']

    # Simple POS approximation by suffix patterns
    nouns   = [w for w in words if w.endswith(('tion','ness','ment','ity','ism','ist'))]
    verbs   = [w for w in words if w.endswith(('ed','ing','ize','ise','ate','fy'))]
    adjectives = [w for w in words if w.endswith(('ful','less','ous','ive','able','ible','al','ic'))]
    adverbs = [w for w in words if w.endswith('ly')]

    # Named entity approximation (capitalized words not at sentence start)
    ner_candidates = []
    for sent in sentences:
        sent_words = sent.split()
        for i, w in enumerate(sent_words[1:], 1):
            if w and w[0].isupper() and len(w) > 2:
                ner_candidates.append(w)

    # Passive voice detection
    passive_pattern = re.compile(r'\b(is|are|was|were|be|been|being)\s+\w+ed\b', re.I)
    passive_count   = len(passive_pattern.findall(text))

    # Transition words
    transitions = ['however','therefore','furthermore','moreover','consequently',
                   'nevertheless','additionally','in conclusion','as a result',
                   'for example','in contrast','on the other hand','similarly',
                   'first','second','third','finally','in summary']
    transition_found = [t for t in transitions if t in text.lower()]

    # Average sentence length
    avg_sent_len = sum(len(s.split()) for s in sentences) / max(len(sentences), 1)

    return {
        'nouns':            nouns[:20],
        'verbs':            verbs[:20],
        'adjectives':       adjectives[:20],
        'adverbs':          adverbs[:20],
        'named_entities':   list(set(ner_candidates))[:15],
        'passive_count':    passive_count,
        'transitions':      transition_found,
        'avg_sent_len':     round(avg_sent_len, 1),
        'pos_ratio': {
            'nouns':      len(nouns),
            'verbs':      len(verbs),
            'adjectives': len(adjectives),
            'adverbs':    len(adverbs)
        }
    }


# ==============================================================
#   STEP 3 — TEN ANALYSIS MODULES
# ==============================================================

# ── 3.1 Grammar Detection ────────────────────────────────────

def check_grammar(text: str, preprocessed: dict) -> dict:
    """Detect common grammar issues."""
    errors   = []
    warnings = []

    # Double spaces
    if '  ' in text:
        warnings.append("Multiple spaces detected between words")

    # Sentence fragments (very short sentences)
    for s in preprocessed['sentences']:
        if len(s.split()) < 4:
            warnings.append(f"Possible sentence fragment: '{s[:50]}'")

    # Missing punctuation check
    paragraphs = text.split('\n')
    for p in paragraphs:
        p = p.strip()
        if p and len(p) > 30 and not p[-1] in '.!?':
            warnings.append("Paragraph may be missing end punctuation")
            break

    # Common confusion pairs
    confused = {
        r'\btheir\s+is\b': "their is → there is",
        r'\bto\s+many\b':  "to many → too many",
        r'\bshouldnt\b':   "shouldnt → shouldn't",
        r'\bcant\b':       "cant → can't",
        r'\bwont\b':       "wont → won't",
    }
    for pattern, fix in confused.items():
        if re.search(pattern, text, re.I):
            errors.append(f"Possible error: {fix}")

    score = max(0, 100 - (len(errors) * 15) - (len(warnings) * 5))
    return {
        'score':    min(score, 100),
        'errors':   errors,
        'warnings': warnings,
        'label':    'Grammar & Mechanics'
    }


# ── 3.2 Vocabulary Enhancement ───────────────────────────────

def analyze_vocabulary(preprocessed: dict) -> dict:
    """Check vocabulary richness and suggest upgrades."""
    words       = preprocessed['words']
    word_freq   = preprocessed['word_freq']
    total_words = len(words)

    if total_words == 0:
        return {'score': 0, 'issues': [], 'suggestions': [], 'label': 'Vocabulary'}

    unique_words = len(set(words))
    diversity_ratio = unique_words / total_words

    # Overused words
    overused = [(w, c) for w, c in word_freq.most_common(10) if c > 3]

    # Weak words to strengthen
    weak_words = {
        'good':   ['excellent','outstanding','exceptional','remarkable'],
        'bad':    ['inadequate','detrimental','problematic','suboptimal'],
        'big':    ['substantial','significant','considerable','extensive'],
        'get':    ['obtain','acquire','achieve','secure'],
        'make':   ['develop','establish','create','implement'],
        'show':   ['demonstrate','illustrate','reveal','highlight'],
        'use':    ['utilize','employ','leverage','implement'],
        'help':   ['facilitate','support','assist','enable'],
        'think':  ['believe','consider','assert','contend'],
        'very':   ['exceptionally','remarkably','significantly','notably'],
    }

    weak_found = []
    for weak, alts in weak_words.items():
        count = words.count(weak)
        if count > 0:
            weak_found.append({
                'word':        weak,
                'count':       count,
                'suggestions': alts
            })

    score = int(diversity_ratio * 100)
    score = max(40, min(score, 100))

    return {
        'score':            score,
        'diversity_ratio':  round(diversity_ratio, 2),
        'unique_words':     unique_words,
        'total_words':      total_words,
        'overused':         overused[:5],
        'weak_words':       weak_found[:5],
        'label':            'Vocabulary & Word Choice'
    }


# ── 3.3 Semantic Coherence ───────────────────────────────────

def analyze_coherence(preprocessed: dict, nlp_data: dict) -> dict:
    """Check logical flow and semantic consistency."""
    sentences   = preprocessed['sentences']
    transitions = nlp_data['transitions']
    issues      = []

    # Check transition usage
    transition_ratio = len(transitions) / max(len(sentences), 1)

    if len(transitions) == 0:
        issues.append("No transition words found — add words like 'however', 'therefore', 'furthermore'")
    elif transition_ratio < 0.1:
        issues.append("Low use of transition words — improve paragraph flow")

    # Topic drift: check if first and last sentence share keywords
    if len(sentences) >= 3:
        first_words = set(re.findall(r'\b[a-z]{4,}\b', sentences[0].lower()))
        last_words  = set(re.findall(r'\b[a-z]{4,}\b', sentences[-1].lower()))
        overlap = first_words & last_words
        if len(overlap) < 2:
            issues.append("Opening and closing sentences have low keyword overlap — check topic consistency")

    score = max(40, 100 - (len(issues) * 20) + (len(transitions) * 5))
    score = min(score, 100)

    return {
        'score':      score,
        'transitions': transitions,
        'issues':      issues,
        'label':       'Semantic Coherence'
    }


# ── 3.4 Topic Relevance ──────────────────────────────────────

def check_topic_relevance(text: str, preprocessed: dict) -> dict:
    """Check relevance to scholarship essay topics."""
    scholarship_keywords = [
        'scholarship','academic','education','research','leadership','community',
        'achievement','goal','aspiration','university','study','career','future',
        'contribution','impact','challenge','overcome','passion','dream','opportunity',
        'develop','learn','grow','inspire','dedicate','commit','pursue'
    ]

    text_lower   = text.lower()
    words        = preprocessed['words']
    found_kw     = [kw for kw in scholarship_keywords if kw in text_lower]
    relevance_pct = len(found_kw) / len(scholarship_keywords)

    missing = [kw for kw in scholarship_keywords[:12] if kw not in text_lower][:5]
    issues  = []

    if relevance_pct < 0.3:
        issues.append("Essay lacks scholarship-specific keywords — add more academic/career-focused language")
    if 'future' not in text_lower and 'goal' not in text_lower:
        issues.append("Missing future goals — scholarship essays should discuss your aspirations")

    score = int(relevance_pct * 100)
    score = max(30, min(score, 100))

    return {
        'score':          score,
        'found_keywords': found_kw,
        'missing_kw':     missing,
        'relevance_pct':  round(relevance_pct * 100, 1),
        'issues':         issues,
        'label':          'Topic Relevance'
    }


# ── 3.5 Readability Analyzer ─────────────────────────────────

def analyze_readability(preprocessed: dict) -> dict:
    """Flesch-Kincaid readability score."""
    words       = preprocessed['words']
    sentences   = preprocessed['sentences']
    word_count  = len(words)
    sent_count  = max(len(sentences), 1)

    # Count syllables (approximate)
    def syllable_count(word):
        word  = word.lower()
        count = len(re.findall(r'[aeiou]+', word))
        if word.endswith('e') and count > 1:
            count -= 1
        return max(count, 1)

    total_syllables = sum(syllable_count(w) for w in words)

    if word_count == 0:
        return {'score': 0, 'grade_level': 'N/A', 'label': 'Readability'}

    # Flesch Reading Ease
    fre = 206.835 - (1.015 * (word_count / sent_count)) - (84.6 * (total_syllables / word_count))
    fre = max(0, min(fre, 100))

    # Grade level
    if fre >= 70:   grade = "Easy (High School)"
    elif fre >= 50: grade = "Standard (College)"
    elif fre >= 30: grade = "Difficult (Academic)"
    else:           grade = "Very Complex"

    issues = []
    avg_sent = word_count / sent_count
    if avg_sent > 30:
        issues.append(f"Average sentence length ({avg_sent:.0f} words) is too long — aim for 15-25 words")
    if avg_sent < 8:
        issues.append("Sentences are too short — combine some for better flow")

    return {
        'score':       int(fre),
        'fre_score':   round(fre, 1),
        'grade_level': grade,
        'avg_sent_len': round(avg_sent, 1),
        'issues':      issues,
        'label':       'Readability'
    }


# ── 3.6 Tone Detector ────────────────────────────────────────

def detect_tone(text: str) -> dict:
    """Detect essay tone: formal/informal/positive/negative."""
    text_lower = text.lower()
    words      = re.findall(r'\b[a-z]+\b', text_lower)

    # Informal markers
    informal = ["gonna","wanna","kinda","sorta","gotta","ya","ain't",
                "dunno","stuff","things","lot","lots","really","pretty","super"]
    informal_found = [w for w in informal if w in words]

    # Formal markers
    formal = ["therefore","consequently","furthermore","moreover","nevertheless",
              "subsequently","accordingly","thus","hence","whereby"]
    formal_found = [w for w in formal if w in words]

    # Positive/negative sentiment
    positive = ["achieve","success","excel","improve","grow","lead","inspire",
                "dedicate","passionate","committed","driven","innovative"]
    negative = ["difficult","challenge","struggle","problem","fail","lack","poor"]
    pos_count = sum(1 for w in positive if w in text_lower)
    neg_count = sum(1 for w in negative if w in text_lower)

    formality_score = min(len(formal_found) * 10 - len(informal_found) * 15 + 60, 100)
    formality_score = max(0, formality_score)

    issues = []
    if informal_found:
        issues.append(f"Informal language detected: {', '.join(informal_found[:4])} — replace with formal alternatives")
    if neg_count > pos_count + 2:
        issues.append("Essay tone is too negative — balance challenges with positive outcomes")

    return {
        'score':          formality_score,
        'formal_words':   formal_found,
        'informal_words': informal_found,
        'positive_count': pos_count,
        'negative_count': neg_count,
        'tone':           'Formal' if formality_score > 60 else 'Semi-Formal' if formality_score > 40 else 'Informal',
        'issues':         issues,
        'label':          'Tone & Voice'
    }


# ── 3.7 Redundancy Detector ──────────────────────────────────

def detect_redundancy(preprocessed: dict, text: str) -> dict:
    """Find repeated phrases, redundant expressions."""
    sentences  = preprocessed['sentences']
    word_freq  = preprocessed['word_freq']
    issues     = []
    redundant  = []

    # Find repeated phrases (3+ word phrases)
    words = text.lower().split()
    phrase_freq = Counter()
    for i in range(len(words) - 2):
        phrase = ' '.join(words[i:i+3])
        if not any(sw in phrase for sw in ['the a an is are was were']):
            phrase_freq[phrase] += 1

    repeated_phrases = [(p, c) for p, c in phrase_freq.most_common(10) if c > 1]

    # Redundant expressions
    redundancies = {
        'absolutely essential':     'essential',
        'completely finished':      'finished',
        'future plans':             'plans',
        'past history':             'history',
        'end result':               'result',
        'basic fundamentals':       'fundamentals',
        'true fact':                'fact',
        'advance planning':         'planning',
    }
    found_redundant = []
    for phrase, fix in redundancies.items():
        if phrase in text.lower():
            found_redundant.append(f"'{phrase}' → use just '{fix}'")

    if repeated_phrases:
        issues.append(f"Repeated phrases detected: {len(repeated_phrases)} found")
    if found_redundant:
        issues.extend(found_redundant[:3])

    # Overused words (top content words used 4+ times)
    overused_words = [(w, c) for w, c in word_freq.most_common(8) if c >= 4]

    score = max(40, 100 - (len(issues) * 12) - (len(overused_words) * 5))
    score = min(score, 100)

    return {
        'score':            score,
        'repeated_phrases': repeated_phrases[:5],
        'overused_words':   overused_words[:5],
        'redundancies':     found_redundant[:5],
        'issues':           issues,
        'label':            'Redundancy'
    }


# ── 3.8 Structure Validator ──────────────────────────────────

def validate_structure(text: str, preprocessed: dict) -> dict:
    """Check essay structure: intro, body, conclusion."""
    paragraphs = [p.strip() for p in text.split('\n') if len(p.strip()) > 30]
    sentences  = preprocessed['sentences']
    issues     = []
    suggestions = []

    has_intro      = len(paragraphs) >= 1
    has_body       = len(paragraphs) >= 2
    has_conclusion = len(paragraphs) >= 3

    # Check conclusion markers
    conclusion_markers = ['in conclusion','to conclude','in summary','therefore',
                          'finally','ultimately','in closing','to summarize']
    last_para = paragraphs[-1].lower() if paragraphs else ''
    has_conclusion_marker = any(m in last_para for m in conclusion_markers)

    # Check if introduction has a hook
    intro_questions  = '?' in paragraphs[0] if paragraphs else False
    intro_statistics = bool(re.search(r'\d+%|\d+ percent|\d+ million', paragraphs[0] if paragraphs else ''))
    has_hook = intro_questions or intro_statistics or len(paragraphs[0]) > 100

    if not has_intro:
        issues.append("Missing introduction paragraph")
    if not has_body:
        issues.append("Add at least 2-3 body paragraphs to support your thesis")
        suggestions.append("Add a paragraph about your academic achievements")
        suggestions.append("Add a paragraph about your career goals")
    if not has_conclusion:
        issues.append("Missing conclusion paragraph")
    if not has_conclusion_marker:
        suggestions.append("Start your conclusion with a marker: 'In conclusion', 'Therefore', etc.")
    if not has_hook:
        suggestions.append("Start with a compelling hook: a question, statistic, or personal anecdote")

    para_count = len(paragraphs)
    score = 0
    if has_intro:      score += 25
    if has_body:       score += 35
    if has_conclusion: score += 25
    if has_hook:       score += 15
    score = min(score, 100)

    return {
        'score':                  score,
        'paragraph_count':        para_count,
        'has_intro':              has_intro,
        'has_body':               has_body,
        'has_conclusion':         has_conclusion,
        'has_conclusion_marker':  has_conclusion_marker,
        'has_hook':               has_hook,
        'issues':                 issues,
        'suggestions':            suggestions,
        'label':                  'Structure & Organization'
    }


# ── 3.9 Argument Strength ────────────────────────────────────

def analyze_arguments(text: str, preprocessed: dict) -> dict:
    """Check strength and support of arguments."""
    text_lower  = text.lower()
    sentences   = preprocessed['sentences']
    issues      = []
    suggestions = []

    # Evidence markers
    evidence_words = ['for example','for instance','specifically','such as',
                      'research shows','studies show','according to','data shows',
                      'statistically','evidence suggests','proven','demonstrated']
    evidence_found = [e for e in evidence_words if e in text_lower]

    # Claim words
    claim_words = ['believe','argue','contend','assert','claim','propose',
                   'suggest','maintain','demonstrate','prove']
    claims_found = [c for c in claim_words if c in text_lower]

    # Numbers/statistics (strong evidence)
    numbers = re.findall(r'\b\d+(?:\.\d+)?(?:%|percent|million|billion|thousand)?\b', text)

    if not evidence_found:
        issues.append("No evidence or examples found — add specific examples to support your claims")
        suggestions.append("Add 'For example,...' or 'For instance,...' to support your main points")
    if not numbers:
        suggestions.append("Add statistics or quantifiable achievements to strengthen your argument")
        suggestions.append("E.g., 'I maintained a 3.8 GPA while leading a team of 15 students'")
    if not claims_found:
        suggestions.append("Make your stance clearer — use words like 'I believe', 'I argue', 'I am committed to'")

    score = 40
    score += len(evidence_found) * 8
    score += len(numbers) * 5
    score += len(claims_found) * 4
    score = min(score, 100)

    return {
        'score':          score,
        'evidence_found': evidence_found[:5],
        'claims_found':   claims_found[:5],
        'numbers_found':  numbers[:5],
        'issues':         issues,
        'suggestions':    suggestions,
        'label':          'Argument Strength'
    }


# ── 3.10 Plagiarism Checker ──────────────────────────────────

def check_plagiarism(text: str, preprocessed: dict) -> dict:
    """
    Basic originality check (no external API).
    Detects clichés and overly common phrases.
    """
    text_lower = text.lower()
    cliches    = [
        'since the dawn of time','in today\'s world','in this day and age',
        'it goes without saying','at the end of the day','think outside the box',
        'passion for learning','dedicated to making a difference',
        'always been passionate','from a young age','dream come true',
        'hard work pays off','strive to be the best','make the world a better place',
        'overcome adversity','reach my full potential'
    ]

    found_cliches = [c for c in cliches if c in text_lower]
    issues        = []
    suggestions   = []

    if found_cliches:
        issues.append(f"Clichéd phrases detected ({len(found_cliches)}) — make your writing more original")
        for c in found_cliches[:3]:
            suggestions.append(f"Replace '{c}' with a specific personal story or original phrasing")

    # Originality score based on clichés
    originality_score = max(50, 100 - (len(found_cliches) * 15))

    return {
        'score':           originality_score,
        'cliches_found':   found_cliches,
        'originality_pct': originality_score,
        'issues':          issues,
        'suggestions':     suggestions,
        'label':           'Originality Check'
    }


# ==============================================================
#   STEP 4 — LLM REWRITER (Groq API)
# ==============================================================

import json
import urllib.request

GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# ==============================================================
#   STEP 4 — LLM REWRITER (Groq API)
# ==============================================================

def call_llm_rewriter(original_text: str, analysis_summary: str) -> str:
    """
    Call Groq API to rewrite the essay based on analysis findings.
    """
    global GROQ_API_KEY
    
    if not GROQ_API_KEY:
        print("⚠️ GROQ_API_KEY not set")
        return None

    # ✅ FIX: Use the same model as before
    models_to_try = ["mixtral-8x7b-32768", "llama3-70b-8192", "llama-3.3-70b-versatile"]
    
    system = """You are an expert scholarship essay editor and writing coach.
Your task is to rewrite the given essay to make it:
- More compelling and authentic
- Academically formal but personal
- Well-structured with clear intro, body, and conclusion
- Rich with specific examples and strong arguments
- Free of clichés and weak language
- Optimized for scholarship committee readers

Keep the writer's voice and core message intact.
Output ONLY the improved essay — no commentary, no explanation."""

    prompt = f"""ORIGINAL ESSAY:
{original_text}

ISSUES TO FIX:
{analysis_summary}

Please rewrite this essay addressing all the issues above.
Make it compelling, specific, and scholarship-worthy.
Maintain approximately the same length (within ±20%)."""

    for model in models_to_try:
        try:
            print(f"🟡 Trying Groq model: {model}")
            
            # ✅ Use groq library instead of urllib
            from groq import Groq
            client = Groq(api_key=GROQ_API_KEY)
            
            completion = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1500
            )
            
            result = completion.choices[0].message.content
            if result and len(result) > 50:
                print(f"✅ Groq LLM rewriter success with model: {model}")
                return result
            else:
                print(f"⚠️ Groq returned empty response with {model}")
                continue
                
        except Exception as e:
            print(f"⚠️ Groq LLM rewriter failed with {model}: {e}")
            continue
    
    print("❌ All Groq models failed")
    return None


def fallback_rewriter(text: str, analysis: dict) -> str:
    """
    Rule-based rewriter when LLM is unavailable.
    Applies basic improvements to the text.
    """
    improved = text

    # Replace weak words
    replacements = {
        r'\bgood\b':       'excellent',
        r'\bbad\b':        'challenging',
        r'\bvery good\b':  'exceptional',
        r'\bget\b':        'obtain',
        r'\bshow\b':       'demonstrate',
        r'\bhelp\b':       'facilitate',
        r'\bbig\b':        'substantial',
        r'\buse\b':        'utilize',
        r'\bthink\b':      'believe',
        r'\bvery\b':       'significantly',
        r'\breally\b':     'truly',
        r'\blot of\b':     'numerous',
        r'\bthings\b':     'aspects',
        r'\bstuff\b':      'elements',
    }

    for pattern, replacement in replacements.items():
        improved = re.sub(pattern, replacement, improved, flags=re.I)

    # Add conclusion marker if missing
    conclusion_markers = ['in conclusion', 'to conclude', 'in summary', 'therefore', 'finally', 'ultimately']
    last_200_chars = improved[-200:].lower() if len(improved) > 200 else improved.lower()
    
    if not any(m in last_200_chars for m in conclusion_markers):
        # Extract last sentence for conclusion
        sentences = re.split(r'[.!?]+', improved)
        last_sentence = sentences[-2].strip() if len(sentences) >= 2 else "my academic and professional goals"
        improved = improved.rstrip() + f'\n\nIn conclusion, {last_sentence.lower()}.'

    return improved.strip()

# ==============================================================
#   STEP 5 — SCORING ENGINE
# ==============================================================

def calculate_final_score(modules: dict) -> dict:
    """
    Multi-factor weighted scoring engine.

    Weights:
        Grammar         15%
        Vocabulary      10%
        Coherence       10%
        Relevance       15%
        Readability     10%
        Tone            10%
        Redundancy       5%
        Structure       15%
        Arguments       10%
        Originality      0% (bonus)
    """
    weights = {
        'grammar':      0.15,
        'vocabulary':   0.10,
        'coherence':    0.10,
        'relevance':    0.15,
        'readability':  0.10,
        'tone':         0.10,
        'redundancy':   0.05,
        'structure':    0.15,
        'arguments':    0.10,
        'plagiarism':   0.00,   # bonus only
    }

    weighted_sum = 0
    for key, weight in weights.items():
        if key in modules:
            weighted_sum += modules[key]['score'] * weight

    # Bonus for originality
    if 'plagiarism' in modules:
        bonus = (modules['plagiarism']['score'] - 50) * 0.05
        weighted_sum += max(0, bonus)

    total = round(weighted_sum, 1)

    if total >= 85:   grade = 'A'
    elif total >= 75: grade = 'B'
    elif total >= 60: grade = 'C'
    elif total >= 45: grade = 'D'
    else:             grade = 'F'

    return {
        'total':     total,
        'grade':     grade,
        'breakdown': {k: v['score'] for k, v in modules.items()},
        'weights':   weights
    }


# ==============================================================
#   STEP 6 — FEEDBACK GENERATOR
# ==============================================================

def generate_feedback(modules: dict, score: dict) -> dict:
    """
    Generate explainable, actionable feedback from all modules.

    Returns:
        strengths, critical_issues, suggestions, highlights
    """
    strengths        = []
    critical_issues  = []
    all_suggestions  = []
    highlighted      = []

    for module_name, data in modules.items():
        s = data.get('score', 0)

        # Strengths (score > 75)
        if s >= 75:
            strengths.append(f"✅ {data['label']}: {s}/100 — Well done!")

        # Critical issues (score < 50)
        elif s < 50:
            critical_issues.append(f"⚠️ {data['label']}: {s}/100 — Needs improvement")

        # Collect all issues and suggestions
        for issue in data.get('issues', []):
            if issue not in critical_issues:
                critical_issues.append(f"• {issue}")

        for sug in data.get('suggestions', []):
            all_suggestions.append(sug)

    # Top 5 priority suggestions
    priority = all_suggestions[:8]

    # Grade-based overall feedback
    grade = score['grade']
    if grade == 'A':
        overall = "🏆 Excellent essay! Minor refinements will make it perfect."
    elif grade == 'B':
        overall = "👍 Good essay with strong potential. Address the suggestions below."
    elif grade == 'C':
        overall = "📝 Average essay. Focus on structure, evidence, and vocabulary improvements."
    else:
        overall = "🔧 Essay needs significant improvement. Review all flagged sections carefully."

    return {
        'overall':          overall,
        'grade':            grade,
        'strengths':        strengths[:5],
        'critical_issues':  critical_issues[:8],
        'suggestions':      priority,
        'total_score':      score['total']
    }


# ==============================================================
#   MAIN PIPELINE ROUTE
# ==============================================================

@essay_bp.route('/api/essay/analyze', methods=['POST'])
@login_required
def analyze_essay():
    """
    Full analysis pipeline:
      Preprocessing → NLP → 10 Modules → LLM → Scoring → Feedback

    Expects JSON: { essay: str }
    Returns: Complete analysis with improved essay and scores
    """
    try:
        data = request.get_json()
        text = (data or {}).get('essay', '').strip()

        if not text:
            return jsonify({'error': 'Essay text is required'}), 400
        if len(text) < 50:
            return jsonify({'error': 'Essay is too short — minimum 50 characters'}), 400

        # ── Step 1: Preprocess ────────────────────────────────
        preprocessed = preprocess_text(text)

        # ── Step 2: NLP Engine ────────────────────────────────
        nlp_data = run_nlp_engine(preprocessed, text)

        # ── Step 3: Run all 10 modules ────────────────────────
        modules = {
            'grammar':     check_grammar(text, preprocessed),
            'vocabulary':  analyze_vocabulary(preprocessed),
            'coherence':   analyze_coherence(preprocessed, nlp_data),
            'relevance':   check_topic_relevance(text, preprocessed),
            'readability': analyze_readability(preprocessed),
            'tone':        detect_tone(text),
            'redundancy':  detect_redundancy(preprocessed, text),
            'structure':   validate_structure(text, preprocessed),
            'arguments':   analyze_arguments(text, preprocessed),
            'plagiarism':  check_plagiarism(text, preprocessed),
        }

                # ── Step 4: LLM Rewriter ──────────────────────────────
        # Collect all issues for LLM context
        all_issues = []
        for m in modules.values():
            all_issues.extend(m.get('issues', []))
            all_issues.extend(m.get('suggestions', []))
        analysis_summary = '\n'.join(f"- {i}" for i in all_issues[:10])

        # ✅ Try LLM rewriter first
        improved_essay = call_llm_rewriter(text, analysis_summary)
        
        # ✅ Fallback to rule-based if LLM fails
        if not improved_essay or len(improved_essay) < 50:
            print("⚠️ LLM rewriter failed, using fallback rewriter")
            improved_essay = fallback_rewriter(text, modules)
        
        # ✅ Ensure we always have something
        if not improved_essay or len(improved_essay) < 50:
            improved_essay = text  # Ultimate fallback: return original

        # ── Step 5: Scoring Engine ────────────────────────────
        final_score = calculate_final_score(modules)

        # ── Step 6: Feedback Generator ────────────────────────
        feedback = generate_feedback(modules, final_score)

        return jsonify({
            'success':        True,
            'improved_essay': improved_essay,
            'original_stats': {
                'word_count':   preprocessed['word_count'],
                'sent_count':   preprocessed['sent_count'],
                'char_count':   preprocessed['char_count'],
                'avg_sent_len': round(nlp_data['avg_sent_len'], 1),
                'passive_count': nlp_data['passive_count'],
                'transitions':  nlp_data['transitions']
            },
            'modules':        modules,
            'score':          final_score,
            'feedback':       feedback,
        })

    except Exception as e:
        print(f"Essay analysis error: {e}")
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500


# ==============================================================
#   DOWNLOAD ROUTES
# ==============================================================

@essay_bp.route('/api/essay/download/pdf', methods=['POST'])
@login_required
def download_pdf():
    """Generate and download essay as PDF."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
        from reportlab.lib import colors

        data     = request.get_json() or {}
        essay    = data.get('essay', '')
        title    = data.get('title', 'Scholarship Essay')
        score    = data.get('score', 0)

        buffer = BytesIO()
        doc    = SimpleDocTemplate(buffer, pagesize=letter,
                                   rightMargin=72, leftMargin=72,
                                   topMargin=72, bottomMargin=72)

        styles  = getSampleStyleSheet()
        story   = []

        # Title
        title_style = ParagraphStyle('Title',
            parent    = styles['Title'],
            fontSize  = 20,
            textColor = colors.HexColor('#0d62ff'),
            spaceAfter = 12,
            alignment = TA_CENTER
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.1 * inch))

        # Score badge
        score_style = ParagraphStyle('Score',
            parent    = styles['Normal'],
            fontSize  = 11,
            textColor = colors.HexColor('#64748b'),
            spaceAfter = 20,
            alignment = TA_CENTER
        )
        # Essay body
        body_style = ParagraphStyle('Body',
            parent    = styles['Normal'],
            fontSize  = 12,
            leading   = 20,
            textColor = colors.HexColor('#0a2540'),
            alignment = TA_JUSTIFY,
            spaceAfter = 12
        )
        for para in essay.split('\n'):
            para = para.strip()
            if para:
                story.append(Paragraph(para, body_style))
                story.append(Spacer(1, 0.1 * inch))

        # Footer
        footer_style = ParagraphStyle('Footer',
            parent    = styles['Normal'],
            fontSize  = 9,
            textColor = colors.HexColor('#94a3b8'),
            alignment = TA_CENTER,
            spaceBefore = 20
        )
        story.append(Spacer(1, 0.3 * inch))
        story.append(Paragraph("Generated by AI Scholar Hunt", footer_style))

        doc.build(story)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment = True,
            download_name = 'Scholarship_Essay.pdf',
            mimetype      = 'application/pdf'
        )

    except ImportError:
        return jsonify({'error': 'PDF generation requires reportlab. Run: pip install reportlab'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@essay_bp.route('/api/essay/download/docx', methods=['POST'])
@login_required
def download_docx():
    """Generate and download essay as DOCX."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        data  = request.get_json() or {}
        essay = data.get('essay', '')
        title = data.get('title', 'Scholarship Essay')
        score = data.get('score', 0)

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.25)
            section.right_margin  = Inches(1.25)

        # Title
        title_para           = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(0x0d, 0x62, 0xff)
            run.font.size      = Pt(20)

        # Score line
        score_para           = doc.add_paragraph(f"AI Quality Score: {score}/100")
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in score_para.runs:
            run.font.size  = Pt(11)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

        doc.add_paragraph()

        # Essay body
        for para in essay.split('\n'):
            para = para.strip()
            if para:
                p           = doc.add_paragraph(para)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'

        # Footer
        doc.add_paragraph()
        footer           = doc.add_paragraph("Generated by AI Scholar Hunt")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size  = Pt(9)
            run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment = True,
            download_name = 'Scholarship_Essay.docx',
            mimetype      = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except ImportError:
        return jsonify({'error': 'DOCX generation requires python-docx. Run: pip install python-docx'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500
